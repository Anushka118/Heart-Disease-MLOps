from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title Page
    document.add_heading('MLOps Project Report', 0)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('End-to-End Heart Disease Classification Pipeline\n')
    run.bold = True
    run.font.size = Pt(16)
    p.add_run('\n\n\n\n')
    p.add_run('Course: MLOps (S1-25_AIMLCZG523)\n')
    p.add_run('Deliverable: Final End-to-End Pipeline Report\n\n')
    p.add_run('Submitted by: Aalekh\n')
    document.add_page_break()

    # Introduction
    document.add_heading('1. Project Overview', level=1)
    document.add_paragraph(
        "In the modern healthcare landscape, machine learning offers rapid diagnostic capabilities that can assist "
        "medical professionals. However, a standalone model is insufficient for clinical deployment. This project "
        "implements a robust, production-ready Machine Learning Operations (MLOps) pipeline designed for the "
        "classification of heart disease. The primary objective is to automate the entire lifecycle of the machine "
        "learning model, effectively bridging the gap between data science experimentation and production operations."
    )
    document.add_paragraph(
        "Key goals achieved in this project include:"
    )
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Reproducibility: Ensuring that data processing and model training can be repeated with consistent results.')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Automation: Reducing manual intervention through CI/CD pipelines and scripted setups.')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Scalability: Utilizing containerization and orchestration to handle varying loads.')
    p = document.add_paragraph(style='List Bullet')
    p.add_run('Observability: Implementing logging and metrics to monitor system health in real-time.')

    # Project Structure
    document.add_heading('2. Repository Structure', level=1)
    document.add_paragraph(
        "The project follows a standard MLOps directory structure to ensure separation of concerns. "
        "Below is a detailed breakdown of the files and their purposes:"
    )
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File / Directory'
    hdr_cells[1].text = 'Description'

    files_data = [
        ('app.py', 'Flask REST API application that serves the model predictions.'),
        ('Dockerfile', 'Defines the production container environment (Python 3.9 slim).'),
        ('k8s_deployment.yaml', 'Kubernetes Deployment definition managing replicas and resources.'),
        ('k8s_service.yaml', 'Kubernetes Service definition exposing the API via LoadBalancer.'),
        ('assignment.ipynb', 'Jupyter notebook for EDA, feature engineering, and model training.'),
        ('heart_disease_pipeline.pkl', 'Serialized Scikit-learn Pipeline (preprocessing + model).'),
        ('scaler.pkl', 'Serialized scaler artifact.'),
        ('requirements.txt', 'Pinned Python dependencies for reproducible environments.'),
        ('test/', 'Directory containing automated test suites (test_model.py).'),
        ('scripts/', 'Helper scripts (e.g., verify_setup.sh).'),
        ('.github/workflows/', 'GitHub Actions CI pipeline definitions.')
    ]

    for filename, desc in files_data:
        row_cells = table.add_row().cells
        row_cells[0].text = filename
        row_cells[1].text = desc

    document.add_paragraph("\n")

    # Data Science Workflow
    document.add_heading('3. Data Science Workflow', level=1)
    
    document.add_heading('3.1 Data Source & Preprocessing', level=2)
    document.add_paragraph(
        "The project utilizes the UCI Heart Disease dataset (ID 45). Data is fetched programmatically using the "
        "`ucimlrepo` library to ensure version consistency vs manual downloads."
    )
    document.add_paragraph("Key Preprocessing Implementation:")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Imputation: Missing values in 'ca' (major vessels) and 'thal' are imputed with the mode (most frequent value) to preserve data volume.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Encoding: Categorical variables ('sex', 'cp', 'slope', etc.) are transformed using One-Hot Encoding to prevent ordinal assumptions.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Scaling: Numerical features are normalized using StandardScaler to unit variance, crucial for the distance-based stability of the model.")

    document.add_heading('3.2 Model Architecture', level=2)
    document.add_paragraph(
        "The final model is encapsulated in a Scikit-learn `Pipeline`. This ensures that raw input data during inference "
        "undergoes the exact same transformations (Scaling, Encoding) as the training data."
    )
    document.add_paragraph("Artifact: `heart_disease_pipeline.pkl` (contains Preprocessor + Classifier).")

    document.add_heading('3.3 Experiment Tracking', level=2)
    document.add_paragraph(
        "MLflow is integrated to track model iterations. It automatically logs:"
    )
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Hyperparameters: max_depth, n_estimators, etc.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Metrics: Accuracy, Recall, F1-score.")
    
    # Engineering & Deployment
    document.add_heading('4. Engineering & Deployment', level=1)
    
    document.add_heading('4.1 API Service (app.py)', level=2)
    document.add_paragraph(
        "The model is served as a microservice using Flask."
    )
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Endpoint: POST /predict")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Input: JSON object representing patient features.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Output: JSON with prediction (0/1) and confidence score.")

    document.add_heading('4.2 Containerization', level=2)
    document.add_paragraph(
        "We use Docker to package the application. The `Dockerfile` uses `python:3.9-slim` to minimize image size. "
        "The serialized model is copied directly into the image, making the container self-sufficient."
    )

    document.add_heading('4.3 Orchestration (Kubernetes)', level=2)
    document.add_paragraph(
        "The application is designed for Kubernetes (Minikube verified).")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Deployment: Manages pod replicas and enforces resource limits (CPU/Memory).")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Service: Exposes the pods via a LoadBalancer on port 80, routing to container port 5000.")

    # CI/CD
    document.add_heading('5. Continuous Integration (CI/CD)', level=1)
    document.add_paragraph(
        "Continuous Integration is managed via GitHub Actions (`.github/workflows/mlops_pipeline.yml`). "
        "Every push to `main` triggers:"
    )
    document.add_paragraph("1. Environment Setup (Python 3.9)", style='List Number')
    document.add_paragraph("2. Linting (flake8): Checks for syntax errors and coding standards.", style='List Number')
    document.add_paragraph("3. Automated Testing (pytest): Executing `test/test_model.py`.", style='List Number')
    
    document.add_paragraph("The test suite includes:")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Existence Checks: Verifying model artifacts (`.pkl`) and data files exist.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Model Loading: ensuring the pipeline object deserializes correctly.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Inference Smoke Test: Running a dummy prediction to verify output shape and type.")

    # Monitoring
    document.add_heading('6. Monitoring & Observability', level=1)
    document.add_paragraph(
        "The application is instrumented with `prometheus-flask-exporter`. "
        "The `/metrics` endpoint exposes:"
    )
    p = document.add_paragraph(style='List Bullet')
    p.add_run("System Metrics: CPU usage, Memory usage.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Application Metrics: Request latency, Status codes.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Business Metrics: `prediction_requests` counter to track model usage volume.")

    # Conclusion
    document.add_heading('7. Conclusion', level=1)
    document.add_paragraph(
        "This project successfully demonstrates the transition of a machine learning model from a Jupyter notebook "
        "to a production-grade microservice. By establishing pipelines for Data, Code, and Deployment, we minimize "
        "manual errors and accelerate the time-to-market for model updates."
    )
    
    document.add_page_break()
    document.add_heading('Appendix: Evidence', level=1)
    document.add_paragraph("Refer to the 'screenshots/' folder in the repository for high-resolution evidence.")
    document.add_paragraph("Required screenshots include:")
    document.add_paragraph("1. Deployment Status: Output of `kubectl get all`.")
    document.add_paragraph("2. API Test: `curl` command output showing a successful prediction.")
    document.add_paragraph("3. CI Pipeline: GitHub Actions dashboard showing success.")
    document.add_paragraph("4. Metrics: Browser view of the `/metrics` endpoint.")

    document.save('Heart_Disease_MLOps_Report.docx')
    print("Report generated successfully: Heart_Disease_MLOps_Report.docx")

if __name__ == "__main__":
    create_report()
